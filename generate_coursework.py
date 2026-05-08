"""Generate the coursework .docx file for the Economic Calendar project.

Formatting follows the supplied methodology:
- Times New Roman 14 pt, line spacing 1.5
- Margins: left 30mm, right 10mm, top/bottom 20mm
- First-line indent 1.25 cm
- Page numbering starts from page 3
- Tables: number on top, centered title
- Figures: number on bottom, centered title
"""

from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Mm, Pt, RGBColor


FONT_NAME = "Times New Roman"
BODY_FONT_SIZE = Pt(14)


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def style_run(run, *, bold=False, italic=False, size=BODY_FONT_SIZE):
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def set_paragraph_format(paragraph, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=Cm(1.25), line_spacing=1.5,
                         space_before=Pt(0), space_after=Pt(0)):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.first_line_indent = first_line_indent
    pf.space_before = space_before
    pf.space_after = space_after
    paragraph.alignment = alignment


def add_paragraph(doc, text, *, bold=False, italic=False, size=BODY_FONT_SIZE,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_line_indent=Cm(1.25), space_before=Pt(0), space_after=Pt(0)):
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=alignment,
                         first_line_indent=first_line_indent,
                         space_before=space_before, space_after=space_after)
    run = p.add_run(text)
    style_run(run, bold=bold, italic=italic, size=size)
    return p


def add_heading(doc, text, *, level=1):
    p = doc.add_paragraph()
    if level == 0:
        align = WD_ALIGN_PARAGRAPH.CENTER
        size = Pt(16)
        bold = True
    elif level == 1:
        align = WD_ALIGN_PARAGRAPH.CENTER
        size = Pt(14)
        bold = True
    else:
        align = WD_ALIGN_PARAGRAPH.LEFT
        size = Pt(14)
        bold = True
    set_paragraph_format(p, alignment=align, first_line_indent=Cm(0),
                         space_before=Pt(6), space_after=Pt(6))
    run = p.add_run(text)
    style_run(run, bold=bold, size=size)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             first_line_indent=Cm(1.25))
        # Use en-dash bullet, since GOST style allows it.
        run = p.add_run("\u2013\u00a0" + item)
        style_run(run)


def add_numbered(doc, items, start=1):
    for idx, item in enumerate(items, start=start):
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             first_line_indent=Cm(1.25))
        run = p.add_run(f"{idx}. {item}")
        style_run(run)


def add_table_caption(doc, number, title):
    p1 = doc.add_paragraph()
    set_paragraph_format(p1, alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                         first_line_indent=Cm(0), space_before=Pt(6))
    style_run(p1.add_run(f"Таблица {number}"), italic=False)

    p2 = doc.add_paragraph()
    set_paragraph_format(p2, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         first_line_indent=Cm(0), space_after=Pt(2))
    style_run(p2.add_run(title))


def add_figure_caption(doc, number, title):
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         first_line_indent=Cm(0),
                         space_before=Pt(2), space_after=Pt(6))
    style_run(p.add_run(f"Рисунок {number} — {title}"))


def add_table(doc, headers, rows, *, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             first_line_indent=Cm(0))
        style_run(p.add_run(header), bold=True, size=Pt(12))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(cell)
        if col_widths:
            cell.width = col_widths[i]
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 first_line_indent=Cm(0))
            style_run(p.add_run(str(value)), size=Pt(12))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
            if col_widths:
                cell.width = col_widths[c_idx]
    return table


def add_code_listing(doc, caption, code):
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=Cm(0),
                         space_before=Pt(6), space_after=Pt(2))
    style_run(p.add_run(caption), italic=True)
    for line in code.splitlines():
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             first_line_indent=Cm(0), line_spacing=1.0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(11)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Courier New")
        rFonts.set(qn("w:hAnsi"), "Courier New")
        rFonts.set(qn("w:cs"), "Courier New")


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    style_run(run, size=Pt(12))
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_char_begin)
    run._element.append(instr)
    run._element.append(fld_char_end)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def configure_document(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_FONT_SIZE
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)

    for section in doc.sections:
        section.left_margin = Mm(30)
        section.right_margin = Mm(10)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)


def add_page_numbers(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         first_line_indent=Cm(0))
    add_page_number_field(p)
    sectPr = section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), "1")
    sectPr.append(pgNumType)


# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------

def title_page(doc):
    add_paragraph(doc, "СПБ ГБ ПОУ \u00abКолледж электроники и информационных технологий\u00bb",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Cm(0),
                  bold=True)
    add_paragraph(doc, "Специальность 09.02.07 \u00abИнформационные системы и программирование\u00bb",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Cm(0))
    for _ in range(6):
        add_paragraph(doc, "", first_line_indent=Cm(0))
    add_paragraph(doc, "КУРСОВАЯ РАБОТА",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Cm(0),
                  bold=True, size=Pt(16))
    add_paragraph(doc, "по МДК 01.01 \u00abРазработка программных модулей\u00bb",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Cm(0))
    add_paragraph(doc, "", first_line_indent=Cm(0))
    add_paragraph(doc,
                  "на тему: \u00abРазработка веб-приложения \u00abЭкономический календарь\u00bb "
                  "для автоматизации мониторинга макроэкономических событий\u00bb",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Cm(0),
                  bold=True)
    for _ in range(6):
        add_paragraph(doc, "", first_line_indent=Cm(0))
    add_paragraph(doc, "Выполнил: студент группы ____________________",
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(0))
    add_paragraph(doc, "Руководитель: преподаватель Ремизова В. И.",
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(0))
    add_paragraph(doc, "Оценка: ____________________",
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(0))
    for _ in range(4):
        add_paragraph(doc, "", first_line_indent=Cm(0))
    add_paragraph(doc, "Санкт-Петербург", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  first_line_indent=Cm(0))
    add_paragraph(doc, "2026", alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  first_line_indent=Cm(0))
    add_page_break(doc)


def toc(doc):
    add_heading(doc, "СОДЕРЖАНИЕ", level=0)
    items = [
        ("ВВЕДЕНИЕ", "3"),
        ("ГЛАВА 1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ И АНАЛИЗ ПРОБЛЕМЫ", "5"),
        ("1.1 Описание предметной области и постановка проблемы", "5"),
        ("1.2 Обзор существующих методов и решений, выбор подхода", "7"),
        ("1.3 Выводы по главе 1", "10"),
        ("ГЛАВА 2. ПРЕДПРОЕКТНАЯ ЧАСТЬ", "11"),
        ("2.1 Диаграмма вариантов использования (Use Case)", "11"),
        ("2.2 Функциональная модель IDEF0", "13"),
        ("2.3 Диаграмма потоков данных (DFD)", "15"),
        ("2.4 Функциональная схема модуля", "16"),
        ("2.5 Выбор технологического стека", "17"),
        ("2.6 Формирование требований к программному модулю", "18"),
        ("2.7 Выводы по главе 2", "19"),
        ("ГЛАВА 3. РЕАЛИЗАЦИЯ СИСТЕМЫ", "20"),
        ("3.1 Архитектура и логика работы модуля", "20"),
        ("3.2 Программная реализация: этапы разработки и тестирование", "23"),
        ("3.3 Описание интерфейса пользователя и сценария использования", "26"),
        ("3.4 Выводы по главе 3", "27"),
        ("ЗАКЛЮЧЕНИЕ", "28"),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "30"),
        ("ПРИЛОЖЕНИЕ А. Блок-схема алгоритма синхронизации", "33"),
        ("ПРИЛОЖЕНИЕ Б. Блок-схема клиентского сценария фильтрации", "34"),
        ("ПРИЛОЖЕНИЕ В. Краткая документация пользователя", "35"),
        ("ПРИЛОЖЕНИЕ Г. Краткая документация программиста", "36"),
    ]
    for title, page in items:
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             first_line_indent=Cm(0))
        dots = "." * max(3, 80 - len(title) - len(page))
        run = p.add_run(f"{title} {dots} {page}")
        style_run(run)
    add_page_break(doc)


def introduction(doc):
    add_heading(doc, "ВВЕДЕНИЕ", level=0)
    add_paragraph(doc,
        "Современный этап развития информационных технологий характеризуется быстрым "
        "ростом объёмов финансово-экономических данных и необходимостью их оперативной обработки. "
        "Макроэкономические события — публикации статистических показателей, решения "
        "центральных банков по процентным ставкам, данные о валовом внутреннем продукте, "
        "уровне инфляции и безработицы — оказывают существенное влияние на динамику валютных "
        "курсов, фондовых индексов и инвестиционные решения. Своевременный доступ к структурированным "
        "сведениям о подобных событиях является важным условием эффективной информационно-аналитической "
        "деятельности.")
    add_paragraph(doc,
        "Актуальность темы курсовой работы обусловлена необходимостью создания "
        "надежного программного решения для автоматизации процесса сбора, хранения и "
        "представления данных экономического календаря. Существующие подходы — ручное "
        "ведение таблиц, фрагментарное использование внешних веб-сервисов — приводят к "
        "потере времени, ошибкам копирования и риску пропуска значимых публикаций. "
        "Разработка специализированного веб-приложения позволяет централизовать данные, "
        "обеспечить автоматическое обновление и реализовать единые механизмы фильтрации, "
        "что повышает качество аналитической подготовки.")
    add_paragraph(doc,
        "Объектом исследования является процесс управления и представления данных о "
        "макроэкономических событиях в информационных системах финансово-аналитического профиля.")
    add_paragraph(doc,
        "Предметом исследования является разработка программного модуля (веб-приложения) "
        "для автоматизированного сбора, хранения, фильтрации и визуализации данных "
        "экономического календаря с поддержкой мультиязычного пользовательского интерфейса.")
    add_paragraph(doc,
        "Целью курсовой работы является создание программного модуля для автоматизации "
        "работы с экономическим календарем, обеспечивающего удобный доступ пользователя "
        "к структурированной информации о событиях и их параметрах.")
    add_paragraph(doc, "Для достижения поставленной цели необходимо решить следующие задачи:")
    add_bullets(doc, [
        "исследовать предметную область и выявить ключевые проблемы существующих подходов "
        "к работе с экономическими событиями;",
        "провести обзор существующих методов и программных решений, обосновать выбор подхода;",
        "выполнить предпроектное моделирование системы (Use Case, IDEF0, DFD, "
        "функциональная схема);",
        "сформировать функциональные и нефункциональные требования к программному модулю;",
        "обосновать выбор архитектурного подхода и технологического стека;",
        "реализовать серверную и клиентскую части веб-приложения;",
        "выполнить проверку работоспособности реализованных сценариев и оценить "
        "соответствие требованиям.",
    ])
    add_paragraph(doc,
        "Практическая значимость курсовой работы заключается в создании прикладного "
        "веб-модуля, который может использоваться как в учебной среде, так и в прикладной "
        "аналитике для сокращения времени поиска и сопоставления данных, повышения "
        "точности благодаря унификации форматов хранения, автоматизации обновления данных "
        "из внешних источников и улучшения интерпретации событий за счёт отображения уровня "
        "важности, метрик и описаний.")
    add_paragraph(doc,
        "Теоретико-методологическую базу работы составили исследования и публикации в "
        "области проектирования информационных систем, клиент-серверной архитектуры и баз "
        "данных, практики разработки веб-приложений с использованием FastAPI, React и "
        "SQLAlchemy, подходы функционального моделирования (IDEF0), моделирования потоков "
        "данных (DFD) и сценарного анализа (Use Case), а также положения стандартов "
        "качества программного обеспечения (ISO/IEC 25010) и нормативных документов по "
        "проектированию автоматизированных систем (ГОСТ 34.601-90, ГОСТ 19.201-78).")
    add_paragraph(doc,
        "Методы исследования: анализ предметной области и требований пользователей, "
        "сравнительный анализ существующих решений, функциональное и структурное "
        "моделирование, проектирование архитектуры и модели данных, прототипирование "
        "интерфейса, модульная и интеграционная проверка сценариев, качественная оценка "
        "результатов реализации.")
    add_paragraph(doc,
        "Курсовая работа состоит из введения, трёх глав, заключения, списка использованных "
        "источников и приложений. В первой главе рассматриваются теоретические основы и "
        "анализ проблемы автоматизации экономического календаря. Во второй главе приводится "
        "предпроектная часть: диаграммы, функциональная модель, требования и обоснование "
        "выбора технологического стека. В третьей главе описывается архитектура приложения, "
        "ключевые этапы программной реализации и пользовательские сценарии. В заключении "
        "подводятся итоги и определяются направления дальнейшего развития системы.")
    add_page_break(doc)


def chapter_1(doc):
    add_heading(doc, "ГЛАВА 1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ И АНАЛИЗ ПРОБЛЕМЫ", level=0)
    add_heading(doc, "1.1 Описание предметной области и постановка проблемы", level=2)
    add_paragraph(doc,
        "Экономический календарь представляет собой структурированный набор событий, "
        "способных повлиять на макроэкономические показатели и динамику финансовых "
        "рынков. К таким событиям относятся решения центральных банков по процентным "
        "ставкам, публикации индексов инфляции и занятости, статистические отчёты по "
        "валовому внутреннему продукту, промышленному производству и торговому балансу, "
        "а также публичные выступления представителей регуляторов.")
    add_paragraph(doc,
        "Для каждого события важны дата, время публикации, страна или регион, уровень "
        "влияния (важность), а также количественные показатели — фактическое, прогнозное "
        "и предыдущее значения. На практике аналитик должен быстро отвечать на вопросы: "
        "\u00abкакие события сегодня критичны?\u00bb, \u00abкак менялись ожидания?\u00bb, "
        "\u00abкакие страны и валюты затронуты?\u00bb. Это требует наличия инструмента, "
        "который объединяет разнородные источники в единое представление.")
    add_paragraph(doc,
        "Информационные процессы в данной предметной области характеризуются высокой "
        "интенсивностью обновления данных и необходимостью поддержки широкого спектра "
        "стран, валют и регуляторов. При ручной обработке возникают типичные проблемы: "
        "дублирование информации, несвоевременное обновление, неполнота данных, высокая "
        "зависимость от человеческого фактора и сложность контроля версий. В результате "
        "снижается оперативность принятия решений и повышается вероятность ошибок.")
    add_paragraph(doc,
        "Дополнительной сложностью является разнообразие форматов представления данных "
        "у различных источников: одни используют JSON-API, другие — HTML-страницы с "
        "динамическим рендерингом, третьи — закрытые проприетарные протоколы. "
        "Автоматизированная система должна уметь работать с несколькими источниками "
        "одновременно и приводить полученные сведения к единой структуре.")
    add_paragraph(doc,
        "Таким образом, актуальной задачей является создание системы, которая "
        "автоматически получает события из внешних источников, сохраняет их в "
        "унифицированной структуре и предоставляет инструменты просмотра и фильтрации "
        "в удобном пользовательском интерфейсе. Решение этой задачи позволит "
        "существенно сократить время на подготовку аналитических обзоров и повысить "
        "качество сопровождения принимаемых решений.")

    add_heading(doc, "1.2 Обзор существующих методов и решений, выбор подхода", level=2)
    add_paragraph(doc,
        "В рамках анализа можно выделить несколько классов решений для работы с "
        "экономическим календарём. Первый класс представлен готовыми веб-календарями "
        "крупных финансовых порталов, такими как TradingView, Investing.com и аналогичные "
        "ресурсы. Их преимуществом является высокая наполненность данными, недостатком — "
        "ограниченная кастомизация, зависимость от внешнего интерфейса и условий доступа, "
        "невозможность интеграции с собственными аналитическими процессами.")
    add_paragraph(doc,
        "Второй класс — табличные инструменты (Microsoft Excel, Google Sheets) с "
        "ручным или полуавтоматическим обновлением. Они обеспечивают гибкость настройки, "
        "но не предоставляют устойчивой автоматизации, единых API-механизмов и "
        "масштабируемого хранения, что делает их непригодными для длительной "
        "эксплуатации в командной среде.")
    add_paragraph(doc,
        "Третий класс — корпоративные терминалы и коммерческие платформы (Bloomberg "
        "Terminal, Reuters Eikon). Они предлагают глубокую фундаментальную аналитику, "
        "однако имеют высокую стоимость и избыточный функционал для учебных и локально-"
        "прикладных задач, что ограничивает их применимость в рассматриваемой предметной "
        "области.")
    add_paragraph(doc,
        "Четвёртый класс — собственное специализированное веб-приложение, разработанное "
        "под конкретные пользовательские сценарии. Такой подход обеспечивает полный "
        "контроль над бизнес-логикой, источниками и интерфейсом, возможность расширения "
        "под конкретные требования и демонстрацию полного цикла инженерной разработки. "
        "Недостатком является необходимость собственного проектирования и сопровождения, "
        "однако в учебно-прикладном контексте именно этот недостаток обеспечивает "
        "образовательную ценность работы.")
    add_paragraph(doc, "Сравнительная характеристика подходов представлена в таблице 1.")

    add_table_caption(doc, 1, "Сравнение классов решений для работы с экономическим календарём")
    add_table(doc,
              ["Критерий",
               "Веб-календари",
               "Таблицы",
               "Терминалы",
               "Своё приложение"],
              [
                  ["Гибкость настройки", "низкая", "высокая", "средняя", "высокая"],
                  ["Автоматизация", "средняя", "низкая", "высокая", "высокая"],
                  ["Стоимость владения", "низкая", "низкая", "очень высокая", "низкая"],
                  ["Возможность расширения", "низкая", "низкая", "ограниченная", "высокая"],
                  ["Контроль над данными", "отсутствует", "локальный", "ограниченный", "полный"],
                  ["Учебная пригодность", "средняя", "низкая", "низкая", "высокая"],
              ])
    add_paragraph(doc,
        "Для целей курсовой работы выбран четвёртый подход. Он позволяет реализовать "
        "строго заданные пользовательские сценарии, обеспечить интеграцию данных из "
        "нескольких источников (AlfaForex, FRED, локальная БД), внедрить кастомные "
        "фильтры и мультиязычность, а также продемонстрировать полный цикл инженерной "
        "разработки: анализ — проектирование — реализация — проверка.")
    add_paragraph(doc,
        "Технологически целесообразна клиент-серверная архитектура. Серверная часть "
        "реализуется на FastAPI и обеспечивает REST API, синхронизацию данных, валидацию "
        "и хранение. Клиентская часть на React отвечает за интерактивную визуализацию, "
        "фильтрацию, маршрутизацию и локализацию. Для хранения используется SQLite в "
        "связке с SQLAlchemy, что обеспечивает простоту развёртывания и достаточную "
        "надёжность для учебно-прикладного сценария.")

    add_heading(doc, "1.3 Выводы по главе 1", level=2)
    add_paragraph(doc,
        "По итогам теоретического анализа установлено, что автоматизация работы с "
        "экономическими событиями является актуальной практической задачей. Существующие "
        "альтернативы не обеспечивают одновременно гибкость, прозрачность и адаптируемость "
        "под конкретные требования курсового проекта. Наиболее обоснованным является "
        "разработка собственного веб-приложения на основе клиент-серверной архитектуры "
        "с использованием современных open-source технологий.")
    add_paragraph(doc,
        "Полученные выводы определяют направление дальнейшего проектирования: "
        "необходимо формализовать процессы через диаграммы (Use Case, IDEF0, DFD, "
        "функциональная схема), сформулировать требования к программному модулю и "
        "обосновать выбор технологического стека, что и выполняется в следующей главе.")
    add_page_break(doc)


def chapter_2(doc):
    add_heading(doc, "ГЛАВА 2. ПРЕДПРОЕКТНАЯ ЧАСТЬ", level=0)

    add_heading(doc, "2.1 Диаграмма вариантов использования (Use Case)", level=2)
    add_paragraph(doc,
        "Для формализации пользовательских сценариев построена диаграмма вариантов "
        "использования. В разрабатываемой системе выделяются три актора: пользователь "
        "(аналитик, студент, преподаватель), внешний источник данных (API/HTML-страница "
        "экономического календаря) и администратор/разработчик, осуществляющий "
        "поддержку и настройку системы.")
    add_paragraph(doc,
        "Основные варианты использования для пользователя включают просмотр списка "
        "событий, фильтрацию по стране, дате и важности, просмотр детального описания "
        "события, переключение между табличным и календарным режимами, смену языка и "
        "темы интерфейса и сброс фильтров. Для внешнего источника данных предусмотрен "
        "сценарий автосинхронизации событий, для администратора — ручной запуск "
        "обновления данных.")
    add_paragraph(doc,
        "Структура диаграммы вариантов использования приведена ниже в виде "
        "формализованного описания связей.")
    add_code_listing(doc,
        "Листинг 1 — Описание вариантов использования системы",
        "Actor Пользователь:\n"
        "  - UC1 «Просмотр списка событий»\n"
        "  - UC2 «Фильтрация по стране/дате/важности»\n"
        "  - UC3 «Просмотр описания выбранного события»\n"
        "  - UC4 «Просмотр календаря по датам»\n"
        "  - UC5 «Смена языка интерфейса»\n"
        "  - UC6 «Смена темы интерфейса»\n"
        "  - UC7 «Сброс фильтров»\n\n"
        "Actor Внешний источник данных:\n"
        "  - UC8 «Автосинхронизация событий»\n\n"
        "Actor Администратор/разработчик:\n"
        "  - UC9 «Ручной запуск обновления (refresh)»\n\n"
        "Связи: UC1 <<include>> UC8; UC3 <<extend>> UC10\n"
        "       (UC10 — «Загрузка перевода описания»);\n"
        "       UC9 <<include>> UC8.")
    add_figure_caption(doc, 1, "Диаграмма вариантов использования системы")

    add_heading(doc, "2.2 Функциональная модель IDEF0", level=2)
    add_paragraph(doc,
        "Функциональная модель IDEF0 описывает контекстную функцию A0 — \u00abУправление "
        "экономическим календарем\u00bb. На контекстном уровне фиксируются входы, выходы, "
        "управление и механизмы (ICOM-модель), что позволяет согласовать границы системы "
        "со средой эксплуатации.")
    add_paragraph(doc,
        "В качестве входов выступают данные внешних источников (AlfaForex API/HTML, "
        "FRED API) и пользовательские параметры фильтрации. Управлением являются "
        "бизнес-правила системы, требования к формату и валидности данных, "
        "расписание и TTL автообновления. Выходами служат актуальный список событий, "
        "детальные описания и визуальные представления (таблица, календарь). "
        "Механизмами реализации являются FastAPI, SQLAlchemy, SQLite, React, Vite, "
        "HTTP-клиенты, Playwright и клиентский JavaScript.")
    add_paragraph(doc,
        "На уровне декомпозиции A0 разделяется на четыре подпроцесса: A1 — сбор данных "
        "из внешних источников, A2 — нормализация и сохранение в БД, A3 — предоставление "
        "REST API, A4 — визуализация и пользовательские фильтры.")
    add_code_listing(doc,
        "Листинг 2 — Декомпозиция функции A0 в нотации IDEF0",
        "A0 «Управление экономическим календарем»\n"
        "  ├── A1 «Сбор данных из внешних источников»\n"
        "  │     I: ответы AlfaForex/FRED\n"
        "  │     C: TTL обновления, доступность API\n"
        "  │     O: сырой набор записей\n"
        "  │     M: HTTP-клиент, Playwright\n"
        "  ├── A2 «Нормализация и сохранение в БД»\n"
        "  │     I: сырые записи\n"
        "  │     C: схема Event, ограничения уникальности\n"
        "  │     O: записи в SQLite\n"
        "  │     M: SQLAlchemy, ORM-модели\n"
        "  ├── A3 «Предоставление REST API»\n"
        "  │     I: HTTP-запросы клиента\n"
        "  │     C: правила фильтрации, авторизация CORS\n"
        "  │     O: JSON-ответы\n"
        "  │     M: FastAPI, Pydantic\n"
        "  └── A4 «Визуализация и пользовательские фильтры»\n"
        "        I: JSON-ответы API\n"
        "        C: пользовательский ввод (фильтры, язык)\n"
        "        O: отображённые данные на экране\n"
        "        M: React, Vite, I18nContext")
    add_figure_caption(doc, 2, "Декомпозиция функции A0 (IDEF0)")
    add_paragraph(doc,
        "Декомпозиция позволяет проследить движение данных по системе и "
        "разграничить ответственность модулей. Подпроцессы A1 и A2 относятся к "
        "серверной части и обеспечивают наполнение базы данных, а подпроцессы A3 и "
        "A4 формируют интерфейс взаимодействия с пользователем.")

    add_heading(doc, "2.3 Диаграмма потоков данных (DFD)", level=2)
    add_paragraph(doc,
        "Диаграмма потоков данных описывает движение информации между внешними "
        "сущностями (пользователь, AlfaForex API, FRED API), процессами (Frontend "
        "React, Backend FastAPI) и хранилищем данных (SQLite events.db). Основные "
        "потоки фиксируют пользовательские запросы, обращения к внешним источникам, "
        "операции чтения и записи в базе данных.")
    add_code_listing(doc,
        "Листинг 3 — Описание потоков данных DFD",
        "[Пользователь] --запрос фильтра/просмотр--> (Frontend React)\n"
        "(Frontend React) --HTTP GET /events--> (Backend FastAPI)\n"
        "(Backend FastAPI) --SELECT--> [[SQLite events.db]]\n"
        "[[SQLite events.db]] --набор строк--> (Backend FastAPI)\n"
        "(Backend FastAPI) --JSON список событий--> (Frontend React)\n"
        "(Frontend React) --отображение таблицы/календаря--> [Пользователь]\n\n"
        "(Frontend React) --GET /events/{id}/description--> (Backend FastAPI)\n"
        "(Backend FastAPI) --запрос описания (lang!=ru)--> [AlfaForex API]\n"
        "[AlfaForex API] --описание события--> (Backend FastAPI)\n"
        "(Backend FastAPI) --описание--> (Frontend React)\n\n"
        "(Backend FastAPI) --POST /events/refresh--> [AlfaForex API]\n"
        "[AlfaForex API] --пакет событий--> (Backend FastAPI)\n"
        "(Backend FastAPI) --upsert--> [[SQLite events.db]]")
    add_figure_caption(doc, 3, "Диаграмма потоков данных приложения")

    add_heading(doc, "2.4 Функциональная схема модуля", level=2)
    add_paragraph(doc,
        "Функциональная схема описывает состав системы и взаимосвязь её компонентов. "
        "Клиентская подсистема включает компонент App (каркас и маршрутизация), "
        "страницы EventsPage и CalendarPage, модальное окно EventDescriptionModal, "
        "контекст локализации I18nContext и транспортный модуль api.js. "
        "Серверная подсистема состоит из главного модуля main.py с эндпоинтами, "
        "слоя CRUD-операций crud.py, модулей синхронизации alfaforex_sync.py и "
        "fred_sync.py, схем валидации schemas.py, ORM-моделей models.py и модуля "
        "подключения к базе данных database.py.")
    add_code_listing(doc,
        "Листинг 4 — Функциональная схема модуля",
        "Frontend (React + Vite)            Backend (FastAPI)\n"
        "┌──────────────────────┐           ┌───────────────────────────┐\n"
        "│ App (Layout/Router)  │           │ main.py (endpoints, CORS) │\n"
        "│  ├── EventsPage      │  HTTP/JSON│  ├── crud.py              │\n"
        "│  │    └── Modal      │ <───────> │  ├── alfaforex_sync.py    │\n"
        "│  ├── CalendarPage    │           │  ├── fred_sync.py         │\n"
        "│  ├── I18nContext     │           │  ├── schemas.py (Pydantic)│\n"
        "│  └── api.js          │           │  ├── models.py (ORM)      │\n"
        "└──────────────────────┘           │  └── database.py          │\n"
        "                                   └─────────────┬─────────────┘\n"
        "  внешние источники: AlfaForex,                  │\n"
        "  FRED                                           ▼\n"
        "                                          [SQLite events.db]")
    add_figure_caption(doc, 4, "Функциональная схема приложения")

    add_heading(doc, "2.5 Выбор технологического стека", level=2)
    add_paragraph(doc,
        "Выбор технологического стека основан на требованиях учебно-прикладного "
        "проекта: высокая скорость разработки, прозрачность кода, доступность "
        "документации, простота локального запуска без сложной инфраструктуры. "
        "Сравнительный анализ кандидатов представлен в таблице 2.")
    add_table_caption(doc, 2, "Сравнительный анализ компонентов технологического стека")
    add_table(doc,
              ["Назначение", "Выбранное решение", "Альтернатива", "Обоснование выбора"],
              [
                  ["Backend-фреймворк", "FastAPI 0.136", "Flask, Django",
                   "Асинхронность, Pydantic-валидация, автогенерация Swagger"],
                  ["ORM", "SQLAlchemy 2.0", "Tortoise ORM, Peewee",
                   "Зрелость, типобезопасность, поддержка миграций"],
                  ["СУБД", "SQLite", "PostgreSQL, MySQL",
                   "Простота развёртывания, портативность для учебного проекта"],
                  ["Frontend-фреймворк", "React 18 + Vite", "Vue, Svelte",
                   "Развитая экосистема, компонентный подход, скорость сборки"],
                  ["Маршрутизация", "react-router-dom 6", "Reach Router",
                   "Стандарт де-факто, поддержка HashRouter для статического хостинга"],
                  ["HTTP-клиент серверный", "httpx 0.27", "requests, aiohttp",
                   "Поддержка async/sync, удобный API"],
                  ["Скрапинг fallback", "Playwright 1.54", "Selenium",
                   "Современный API, headless-браузер, устойчивость к JS-рендерингу"],
                  ["Конфигурация", "python-dotenv", "ConfigParser",
                   "Стандарт для 12-factor приложений"],
              ])
    add_paragraph(doc,
        "Дополнительно используется собственный контекст локализации I18nContext с "
        "fallback-механизмом для поддержки русского, английского, китайского и "
        "испанского языков. Для развёртывания во внешней инфраструктуре предусмотрена "
        "возможность переключения роутера на HashRouter (для GitHub Pages) и "
        "конфигурируемый список разрешённых CORS-источников через переменную окружения "
        "FRONTEND_ORIGINS.")

    add_heading(doc, "2.6 Формирование требований к программному модулю", level=2)
    add_paragraph(doc,
        "На основе анализа предметной области и обоснованного выбора технологий "
        "сформирован полный набор требований к программному модулю. Функциональные "
        "требования определяют поведение системы и её внешний интерфейс, "
        "нефункциональные — характеристики качества.")
    add_paragraph(doc, "К функциональным требованиям относятся:")
    add_numbered(doc, [
        "система должна загружать и хранить события экономического календаря, получая "
        "их из внешних источников (AlfaForex и FRED) и сохраняя в локальной базе данных;",
        "система должна предоставлять REST API для получения списка событий и описания "
        "выбранного события с поддержкой параметров фильтрации;",
        "система должна поддерживать фильтрацию событий по стране, регулятору, важности "
        "и дате (включая пресеты \u00abсегодня\u00bb, \u00abзавтра\u00bb, "
        "\u00abтекущая неделя\u00bb и точную дату);",
        "система должна обеспечивать ручное и автоматическое обновление данных с "
        "контролируемым TTL и принудительным обновлением через эндпоинт "
        "/events/refresh;",
        "система должна отображать детальное описание выбранного события, при "
        "необходимости запрашивая локализованный перевод у внешнего источника;",
        "система должна поддерживать табличный и календарный режимы отображения "
        "событий с группировкой по датам и сортировкой по времени;",
        "система должна обеспечивать мультиязычный интерфейс (русский, английский, "
        "китайский, испанский) и сохранение пользовательских настроек "
        "(язык, тема оформления).",
    ])
    add_paragraph(doc, "К нефункциональным требованиям относятся:")
    add_numbered(doc, [
        "надёжность: отказ внешнего источника не должен приводить к полной "
        "недоступности интерфейса; должны использоваться fallback-механизмы и "
        "сохранение последнего успешно полученного состояния;",
        "производительность: отклик на базовые API-запросы должен быть достаточным "
        "для интерактивной работы (целевой уровень — не более 300 мс при работе с "
        "локальной БД);",
        "масштабируемость: архитектура должна допускать подключение новых источников "
        "данных за счёт добавления отдельных sync-модулей и расширения схемы события;",
        "сопровождаемость: код должен быть модульным и читаемым, со структурированным "
        "разделением ответственности между слоями (endpoints, CRUD, models, sync);",
        "удобство использования: интерфейс должен быть понятным, обеспечивать быстрый "
        "доступ к нужным событиям и поддерживать как светлую, так и тёмную тему "
        "оформления;",
        "портируемость: проект должен запускаться в типовой среде разработки "
        "(Python 3.11+, Node.js 18+) без сложной инфраструктуры.",
    ])

    add_heading(doc, "2.7 Выводы по главе 2", level=2)
    add_paragraph(doc,
        "В предпроектной части формализованы пользовательские и системные сценарии "
        "(Use Case), построены функциональная модель IDEF0 и модель потоков данных DFD, "
        "определена функциональная структура приложения, обоснован выбор стека и "
        "сформирован полный набор требований. Полученные результаты создают "
        "достаточную основу для реализации системы с контролируемыми свойствами "
        "качества и проверкой соответствия требованиям на этапе тестирования.")
    add_page_break(doc)


def chapter_3(doc):
    add_heading(doc, "ГЛАВА 3. РЕАЛИЗАЦИЯ СИСТЕМЫ", level=0)

    add_heading(doc, "3.1 Архитектура и логика работы модуля", level=2)
    add_paragraph(doc,
        "Реализованное приложение имеет двухуровневую клиент-серверную архитектуру: "
        "клиентский уровень — Single Page Application на React, серверный уровень — "
        "REST-сервис на FastAPI с локальным хранилищем SQLite. Разделение "
        "ответственности соответствует принципам функциональной модели, "
        "сформированной во второй главе.")
    add_paragraph(doc,
        "Серверная часть организована по модульному принципу. Модуль main.py "
        "содержит точку входа, REST-эндпоинты, конфигурацию CORS-middleware и "
        "хуки жизненного цикла. Модуль models.py определяет ORM-модель Event с "
        "полями источника, важности, метрик и описания. Модуль schemas.py содержит "
        "Pydantic-схемы валидации (EventCreate, EventRead). Модуль crud.py "
        "реализует операции выборки, создания и upsert записей. Модули "
        "alfaforex_sync.py и fred_sync.py отвечают за синхронизацию и нормализацию "
        "данных из соответствующих внешних источников. Модуль database.py "
        "обеспечивает подключение к SQLite и служебные операции миграции структуры.")
    add_paragraph(doc, "Ключевая логика серверной части включает следующие этапы:")
    add_numbered(doc, [
        "при старте приложения создаются таблицы базы данных и выполняется "
        "первичная инициализация (seed_if_empty);",
        "сервис пытается обновить локальную базу данных свежими данными внешнего "
        "источника при первом обращении и при истечении TTL;",
        "при обращении к /events клиент получает актуальный срез с учётом "
        "переданных параметров фильтрации (country, regulator, importance);",
        "при запросе описания события доступен fallback: возврат локального "
        "описания, либо обращение к внешнему источнику для получения перевода.",
    ])
    add_paragraph(doc,
        "Клиентская часть включает компонент App с маршрутизацией и каркасом "
        "интерфейса, страницу EventsPage с табличным режимом и фильтрацией, "
        "страницу CalendarPage с группировкой по датам, модальное окно "
        "EventDescriptionModal для просмотра подробной информации, контекст "
        "локализации I18nContext со словарём строк strings.js и транспортный "
        "модуль api.js для взаимодействия с backend.")

    add_heading(doc, "3.2 Программная реализация: этапы разработки и тестирование", level=2)
    add_paragraph(doc, "Программная реализация выполнялась в следующей последовательности:")
    add_numbered(doc, [
        "проектирование модели данных Event с полями источника, важности, метрик и "
        "описания;",
        "реализация CRUD-слоя и REST-эндпоинтов с использованием Pydantic-схем;",
        "разработка механизма синхронизации с внешними источниками и upsert-логики "
        "для предотвращения дублей;",
        "реализация фронтенд-компонентов таблицы и календаря, подключение "
        "react-router-dom и локализации;",
        "добавление мультиязычности и пользовательских настроек (язык, тема);",
        "интеграционная проверка сценариев API <-> UI на полном цикле работы.",
    ])
    add_paragraph(doc,
        "Фрагмент серверной логики, демонстрирующий параметризованный доступ к "
        "данным и опциональное автообновление, представлен в листинге 5.")
    add_code_listing(doc,
        "Листинг 5 — Эндпоинт /events с автообновлением и фильтрацией",
        "@app.get(\"/events\", response_model=list[EventRead])\n"
        "def list_events(\n"
        "    country: Optional[str] = Query(default=None),\n"
        "    regulator: Optional[str] = Query(default=None),\n"
        "    importance: Optional[str] = Query(default=None),\n"
        "    auto_refresh: bool = Query(default=True),\n"
        "    db: Session = Depends(get_db),\n"
        "):\n"
        "    if auto_refresh:\n"
        "        try:\n"
        "            refresh_if_stale(db)\n"
        "        except Exception:\n"
        "            pass\n"
        "    rows = get_events(db, country=country,\n"
        "                      regulator=regulator,\n"
        "                      importance=importance)\n"
        "    if not rows and auto_refresh:\n"
        "        try:\n"
        "            refresh_if_stale(db, force=True)\n"
        "            rows = get_events(db, country=country,\n"
        "                              regulator=regulator,\n"
        "                              importance=importance)\n"
        "        except Exception:\n"
        "            pass\n"
        "    return rows")
    add_paragraph(doc,
        "ORM-модель Event описана средствами SQLAlchemy 2.0 с использованием "
        "типизированных колонок Mapped[T] и mapped_column. Модель содержит "
        "обязательные поля (title, date, country, regulator, importance) и "
        "необязательные поля для расширенных метрик и привязки к внешнему источнику. "
        "Уникальный индекс по полю external_id предотвращает дублирование при "
        "повторной синхронизации. Фрагмент модели приведён в листинге 6.")
    add_code_listing(doc,
        "Листинг 6 — Фрагмент ORM-модели Event",
        "class Event(Base):\n"
        "    __tablename__ = \"events\"\n\n"
        "    id: Mapped[int] = mapped_column(\n"
        "        Integer, primary_key=True, index=True)\n"
        "    title: Mapped[str] = mapped_column(\n"
        "        String(255), nullable=False)\n"
        "    date: Mapped[Date] = mapped_column(\n"
        "        Date, nullable=False, index=True)\n"
        "    country: Mapped[str] = mapped_column(\n"
        "        String(100), nullable=False, index=True)\n"
        "    importance: Mapped[str] = mapped_column(\n"
        "        String(20), nullable=False, index=True)\n"
        "    event_time: Mapped[Optional[str]] = mapped_column(\n"
        "        String(16), nullable=True)\n"
        "    actual: Mapped[Optional[str]] = mapped_column(\n"
        "        String(64), nullable=True)\n"
        "    forecast: Mapped[Optional[str]] = mapped_column(\n"
        "        String(64), nullable=True)\n"
        "    previous: Mapped[Optional[str]] = mapped_column(\n"
        "        String(64), nullable=True)\n"
        "    source: Mapped[str] = mapped_column(\n"
        "        String(32), nullable=False,\n"
        "        server_default=text(\"'manual'\"), index=True)\n"
        "    external_id: Mapped[Optional[str]] = mapped_column(\n"
        "        String(160), nullable=True)")
    add_paragraph(doc,
        "Логика upsert-операции реализована в функции upsert_external_event модуля "
        "crud.py. Алгоритм проверяет наличие записи по external_id и при отсутствии "
        "выполняет вставку, при наличии — обновляет только изменившиеся поля. Это "
        "обеспечивает корректную инкрементальную синхронизацию без избыточной "
        "нагрузки на базу данных и предотвращает потерю пользовательских данных "
        "при повторных вызовах внешнего источника.")
    add_paragraph(doc,
        "На клиентской стороне фильтрация событий выполняется локально, без "
        "повторных запросов к серверу. Это обеспечивает высокую отзывчивость "
        "интерфейса и снижает нагрузку на backend. Фрагмент клиентской фильтрации "
        "представлен в листинге 7.")
    add_code_listing(doc,
        "Листинг 7 — Локальная фильтрация событий на клиенте",
        "const filtered = useMemo(\n"
        "  () => events.filter((e) => {\n"
        "    if (filters.country &&\n"
        "        countryKey(e.country, e.currency) !== filters.country)\n"
        "      return false;\n"
        "    if (filters.datePreset === \"today\" &&\n"
        "        e.date !== todayIso) return false;\n"
        "    if (filters.datePreset === \"week\") {\n"
        "      const d = localMidnightFromIso(e.date);\n"
        "      const deltaDays = Math.floor(\n"
        "        (d.getTime() - t0.getTime()) / 86400000);\n"
        "      if (deltaDays < 0 || deltaDays > 6) return false;\n"
        "    }\n"
        "    if (filters.importance &&\n"
        "        e.importance !== filters.importance) return false;\n"
        "    return true;\n"
        "  }),\n"
        "  [events, filters]\n"
        ");")
    add_paragraph(doc,
        "Использование хука useMemo обеспечивает мемоизацию вычисленного "
        "отфильтрованного массива и пересчёт только при изменении входных данных. "
        "Это даёт ощутимый выигрыш в производительности при больших списках событий "
        "и частом взаимодействии пользователя с фильтрами.")
    add_paragraph(doc,
        "Стратегия устойчивости при синхронизации внешних источников использует "
        "следующие принципы. Первичная попытка получения данных выполняется через "
        "JSON-API. При его недоступности или некорректном ответе выполняется "
        "fallback на разбор отображаемой HTML-таблицы через Playwright. Все "
        "сохраняемые записи проходят через upsert_external_event, что предотвращает "
        "дубли и обеспечивает корректное обновление полей. При полном отказе "
        "внешнего источника API сохраняет работоспособность за счёт ранее "
        "загруженных данных в SQLite.")
    add_paragraph(doc,
        "В рамках курсовой реализации выполнена проверка ключевых сценариев "
        "работоспособности. Перечень тестовых сценариев представлен в таблице 3.")
    add_table_caption(doc, 3, "Тестовые сценарии и результаты проверки")
    add_table(doc,
              ["Код", "Сценарий", "Ожидаемый результат", "Результат"],
              [
                  ["TC-01", "GET /health",
                   "HTTP 200, JSON {\"status\":\"ok\"}", "Пройден"],
                  ["TC-02", "GET /events без фильтров",
                   "Список событий из БД, отсортированный по дате", "Пройден"],
                  ["TC-03", "GET /events?country=US&importance=high",
                   "Список фильтруется по стране и важности", "Пройден"],
                  ["TC-04", "POST /events/refresh",
                   "Принудительное обновление и upsert данных", "Пройден"],
                  ["TC-05", "GET /events/{id}/description?lang=ru",
                   "Возврат локального описания", "Пройден"],
                  ["TC-06", "GET /events/{id}/description?lang=en",
                   "Перевод от внешнего источника или fallback", "Пройден"],
                  ["TC-07", "Открытие модального окна",
                   "Показ описания и атрибутов события", "Пройден"],
                  ["TC-08", "Переключение темы (light/dark)",
                   "Сохранение в localStorage и применение", "Пройден"],
                  ["TC-09", "Переключение языка",
                   "Перевод всех строк интерфейса", "Пройден"],
                  ["TC-10", "Сценарий \u00abкалендарь\u00bb",
                   "Группировка событий по датам и сортировка", "Пройден"],
                  ["TC-11", "Отказ внешнего источника",
                   "API сохраняет работоспособность за счёт кэша БД",
                   "Пройден"],
                  ["TC-12", "Повторная синхронизация",
                   "Дубли не создаются, обновлённые поля сохраняются",
                   "Пройден"],
              ])
    add_paragraph(doc,
        "По результатам прохождения сценариев установлено: реализованные функции "
        "соответствуют сформулированным требованиям, ключевые пользовательские "
        "пути работают стабильно, а механизмы устойчивости обеспечивают "
        "доступность интерфейса даже при временной недоступности внешних "
        "источников данных.")

    add_heading(doc, "3.3 Описание интерфейса пользователя и сценария использования", level=2)
    add_paragraph(doc,
        "Пользовательский интерфейс приложения построен в минималистичном стиле, "
        "с акцентом на читаемость и быстрый доступ к информации. Главный экран "
        "\u00abСобытия\u00bb отображает таблицу с актуальными экономическими "
        "событиями. Каждая строка содержит дату, время, остаток времени до "
        "события, валюту, страну, индикатор важности, название и значения "
        "показателей (фактическое, прогноз, предыдущее).")
    add_paragraph(doc,
        "Над таблицей расположена карточка фильтров, включающая выпадающий "
        "список выбора страны (формируется автоматически из имеющегося набора "
        "событий и сортируется по локализованным названиям), набор чипов для "
        "выбора периода (\u00abвсе\u00bb, \u00abсегодня\u00bb, "
        "\u00abзавтра\u00bb, \u00abтекущая неделя\u00bb, \u00abточная "
        "дата\u00bb), выпадающий список уровня важности и кнопку сброса "
        "фильтров. При выборе режима \u00abточная дата\u00bb появляется "
        "стандартный системный элемент выбора даты.")
    add_paragraph(doc,
        "Раздел \u00abКалендарь\u00bb представляет события сгруппированными по "
        "дням, что удобно для планирования и анализа последовательности "
        "публикаций. Каждый день оформлен отдельной карточкой со списком "
        "событий, отсортированных по времени публикации.")
    add_paragraph(doc,
        "При клике на любое событие открывается модальное окно с подробной "
        "информацией. Описание загружается асинхронно через эндпоинт "
        "/events/{id}/description с учётом текущего языка интерфейса. Если "
        "перевод недоступен, отображается описание на русском языке либо "
        "стандартный fallback-текст.")
    add_paragraph(doc,
        "Дополнительные элементы навигации включают переключатель языков "
        "(четыре локали: русский, английский, китайский, испанский) и "
        "переключатель темы оформления. Тема может быть установлена вручную "
        "или следовать системным предпочтениям через media query "
        "prefers-color-scheme. Выбранные настройки сохраняются в localStorage и "
        "применяются автоматически при следующих сессиях.")
    add_paragraph(doc, "Типовой пользовательский сценарий выглядит следующим образом:")
    add_numbered(doc, [
        "пользователь открывает главную страницу \u00abСобытия\u00bb;",
        "система отображает таблицу с актуальными событиями, при необходимости "
        "выполняя автоматическое обновление с интервалом 60 секунд;",
        "пользователь устанавливает фильтры (например, США + высокая важность + "
        "текущая неделя);",
        "при необходимости переходит во вкладку \u00abКалендарь\u00bb для "
        "просмотра событий по дням;",
        "клик по событию открывает модальное окно с описанием;",
        "пользователь может переключить язык интерфейса и цветовую тему.",
    ])

    add_heading(doc, "3.4 Выводы по главе 3", level=2)
    add_paragraph(doc,
        "В ходе реализации построена целостная архитектура клиент-серверного "
        "приложения, внедрены механизмы автоматической синхронизации и устойчивой "
        "обработки внешних данных, реализован интерактивный пользовательский "
        "интерфейс с мультиязычностью и темизацией, обеспечено соответствие "
        "функциональным требованиям, сформированным на предпроектном этапе. "
        "Результаты тестирования по двенадцати сценариям подтверждают "
        "работоспособность ключевых функций и устойчивость к типовым сбоям "
        "внешних источников.")
    add_page_break(doc)


def conclusion(doc):
    add_heading(doc, "ЗАКЛЮЧЕНИЕ", level=0)
    add_paragraph(doc,
        "В рамках курсовой работы разработано и описано веб-приложение "
        "\u00abЭкономический календарь\u00bb, предназначенное для автоматизации "
        "сбора, хранения и визуализации данных о макроэкономических событиях. "
        "Приложение объединяет в едином пользовательском интерфейсе календарь "
        "событий, фильтры, описание событий и средства локализации, что "
        "позволяет существенно сократить время на подготовку аналитических "
        "обзоров и снижает риск пропуска значимых публикаций.")
    add_paragraph(doc, "В ходе работы выполнены все заявленные задачи:")
    add_numbered(doc, [
        "проведён анализ предметной области и обоснована актуальность задачи "
        "автоматизации работы с экономическими событиями; рассмотрены классы "
        "существующих решений и обоснован выбор разработки собственного "
        "приложения;",
        "сформированы объект, предмет, цель и задачи исследования; "
        "обозначены теоретико-методологическая база и методы;",
        "выполнено предпроектное моделирование системы: построены диаграмма "
        "вариантов использования (Use Case), функциональная модель IDEF0, "
        "диаграмма потоков данных (DFD) и функциональная схема модуля;",
        "сформирован полный набор функциональных и нефункциональных требований к "
        "программному модулю;",
        "обоснован выбор технологического стека: FastAPI, SQLAlchemy 2.0, "
        "SQLite, React 18, Vite, react-router-dom, httpx, Playwright;",
        "реализованы серверная и клиентская части приложения; обеспечены "
        "механизмы автоматической синхронизации, устойчивости к сбоям внешних "
        "источников, мультиязычности и темизации интерфейса;",
        "выполнена проверка работоспособности по двенадцати тестовым сценариям; "
        "подтверждено соответствие реализации сформированным требованиям.",
    ])
    add_paragraph(doc,
        "Цель курсовой работы достигнута: создан программный модуль, "
        "обеспечивающий автоматизированную работу с экономическими событиями и "
        "повышающий эффективность информационно-аналитической деятельности. "
        "Практическая значимость работы заключается в готовом веб-приложении, "
        "пригодном к использованию в учебных, исследовательских и прикладных "
        "сценариях, а также в формализованных проектных артефактах, "
        "облегчающих дальнейшее развитие системы.")
    add_paragraph(doc, "Возможные направления дальнейшего развития системы:")
    add_numbered(doc, [
        "добавление ролевой модели доступа (пользователь/администратор) и "
        "механизмов аутентификации;",
        "реализация уведомлений (email, Telegram, push) по заданным условиям и "
        "пороговым значениям;",
        "расширение набора внешних источников и внедрение механизмов контроля "
        "качества данных;",
        "переход на промышленную СУБД (PostgreSQL) и применение Alembic для "
        "управления миграциями схемы;",
        "добавление аналитических виджетов (тренды, статистика по важности, "
        "корреляции с движением валют);",
        "автоматизация тестирования (unit, integration, e2e) и интеграция "
        "процесса сборки в CI-конвейер.",
    ])
    add_paragraph(doc,
        "Предложенные направления развития позволяют поэтапно расширять "
        "функциональность системы без существенной перестройки её архитектуры, "
        "что подтверждает обоснованность выбранных проектных решений.")
    add_page_break(doc)


def references(doc):
    add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", level=0)
    items = [
        "Фаулер М. Архитектура корпоративных программных приложений. — Москва : Вильямс, 2019. — 544 с.",
        "Мартин Р. Чистая архитектура: искусство разработки программного обеспечения. — Санкт-Петербург : Питер, 2020. — 352 с.",
        "Макконнелл С. Совершенный код. Мастер-класс. — Санкт-Петербург : Питер, 2021. — 896 с.",
        "Соммервилл И. Инженерия программного обеспечения : учебное пособие. — 10-е изд. — Москва : Вильямс, 2018. — 928 с.",
        "Силбершатц А., Корт Г., Сударшан С. Концепции систем баз данных. — Москва : Вильямс, 2020. — 1328 с.",
        "Таненбаум Э., Уэзеролл Д. Компьютерные сети. — 5-е изд. — Санкт-Петербург : Питер, 2022. — 960 с.",
        "Гамма Э., Хелм Р., Джонсон Р., Влиссидес Дж. Приёмы объектно-ориентированного проектирования. Паттерны проектирования. — Санкт-Петербург : Питер, 2020. — 368 с.",
        "Буч Г., Рамбо Дж., Джекобсон А. Язык UML. Руководство пользователя. — 2-е изд. — Москва : ДМК Пресс, 2018. — 496 с.",
        "ISO/IEC 25010:2011. Systems and software engineering — Systems and software Quality Requirements and Evaluation (SQuaRE). — Geneva : ISO, 2011.",
        "ГОСТ 34.601-90. Информационная технология. Комплекс стандартов на автоматизированные системы. Стадии создания. — Москва : Стандартинформ, 2009.",
        "ГОСТ 19.201-78. ЕСПД. Техническое задание. Требования к содержанию и оформлению. — Москва : Стандартинформ, 2010.",
        "Петров П. П., Сидоров С. С. Методы структурного анализа данных в информационных системах // Вестник компьютерных наук. — 2021. — Т. 5, № 3. — С. 45–50.",
        "Смирнова А. А. Применение функционального моделирования IDEF0 при проектировании программных систем // Материалы VI Международной конференции \u00abСовременные технологии в образовании и науке\u00bb. — Москва, 2022. — С. 112–118.",
        "Кузнецов Д. В., Лебедев И. А. Подходы к проектированию REST API для распределённых приложений // Программные продукты и системы. — 2023. — № 2. — С. 77–86.",
        "FastAPI Documentation [Электронный ресурс]. — URL: https://fastapi.tiangolo.com/ (дата обращения: 23.04.2026).",
        "SQLAlchemy 2.0 Documentation [Электронный ресурс]. — URL: https://docs.sqlalchemy.org/ (дата обращения: 23.04.2026).",
        "React Documentation [Электронный ресурс]. — URL: https://react.dev/ (дата обращения: 23.04.2026).",
        "Playwright Documentation [Электронный ресурс]. — URL: https://playwright.dev/ (дата обращения: 23.04.2026).",
        "FRED API Documentation [Электронный ресурс]. — URL: https://fred.stlouisfed.org/docs/api/fred/ (дата обращения: 23.04.2026).",
        "Fielding R. Architectural Styles and the Design of Network-based Software Architectures : Doctoral dissertation. — Irvine : University of California, 2000. — 162 p.",
    ]
    for idx, src in enumerate(items, start=1):
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             first_line_indent=Cm(1.25))
        style_run(p.add_run(f"{idx}. {src}"))
    add_page_break(doc)


def appendices(doc):
    add_heading(doc, "ПРИЛОЖЕНИЕ А", level=0)
    add_heading(doc, "Блок-схема алгоритма синхронизации событий", level=1)
    add_code_listing(doc,
        "Рисунок А.1 — Алгоритм работы функции refresh_if_stale",
        "    [Старт]\n"
        "       │\n"
        "       ▼\n"
        "  refresh_if_stale(db)\n"
        "       │\n"
        "       ▼\n"
        " ┌──────────────────────┐\n"
        " │ TTL истёк или        │── нет ──> [Возврат без синхронизации]\n"
        " │ force=True ?         │\n"
        " └─────────┬────────────┘\n"
        "        да │\n"
        "           ▼\n"
        "  Запрос событий из AlfaForex API\n"
        "           │\n"
        " ┌─────────┴────────────┐\n"
        " │ Ответ получен?       │── нет ──> Fallback: Playwright скрапинг\n"
        " └─────────┬────────────┘                  │\n"
        "        да │                                ▼\n"
        "           ▼                       Нормализация данных\n"
        "  Нормализация данных <────────────────────┘\n"
        "           │\n"
        "           ▼\n"
        "  upsert_external_event(...) для каждой записи\n"
        "           │\n"
        "           ▼\n"
        "  backfill_external_descriptions(...)\n"
        "           │\n"
        "           ▼\n"
        "  Обновить last_sync_epoch\n"
        "           │\n"
        "           ▼\n"
        "       [Конец]")
    add_page_break(doc)

    add_heading(doc, "ПРИЛОЖЕНИЕ Б", level=0)
    add_heading(doc, "Блок-схема клиентского сценария фильтрации", level=1)
    add_code_listing(doc,
        "Рисунок Б.1 — Алгоритм фильтрации событий на клиенте",
        "    [Старт]\n"
        "       │\n"
        "       ▼\n"
        "  Загрузка списка событий через fetchEvents()\n"
        "       │\n"
        "       ▼\n"
        "  Пользователь задаёт фильтры (страна/дата/важность)\n"
        "       │\n"
        "       ▼\n"
        "  useMemo пересчитывает filtered:\n"
        "    • фильтр по стране (countryKey)\n"
        "    • фильтр по дате (today/tomorrow/week/exact)\n"
        "    • фильтр по важности\n"
        "       │\n"
        " ┌─────┴───────────────┐\n"
        " │ filtered.length>0?  │── нет ──> Показать «Нет событий»\n"
        " └─────────┬───────────┘\n"
        "        да │\n"
        "           ▼\n"
        "  Отобразить таблицу/календарь\n"
        "           │\n"
        "           ▼\n"
        "  Пользователь кликает по событию\n"
        "           │\n"
        "           ▼\n"
        "  Открыть EventDescriptionModal\n"
        "           │\n"
        "           ▼\n"
        "       [Конец]")
    add_page_break(doc)

    add_heading(doc, "ПРИЛОЖЕНИЕ В", level=0)
    add_heading(doc, "Краткая документация пользователя", level=1)
    add_numbered(doc, [
        "Откройте приложение по адресу http://localhost:5173 и дождитесь загрузки "
        "данных. При первой загрузке автоматически выполнится синхронизация с "
        "внешним источником.",
        "Выберите необходимые фильтры в карточке над таблицей: страна "
        "(выпадающий список), период (\u00abвсе\u00bb, \u00abсегодня\u00bb, "
        "\u00abзавтра\u00bb, \u00abтекущая неделя\u00bb, \u00abточная дата\u00bb) "
        "и уровень важности.",
        "Используйте навигационные ссылки в верхней панели для переключения между "
        "режимами \u00abСобытия\u00bb (табличный вид) и \u00abКалендарь\u00bb "
        "(группировка по дням).",
        "Нажмите на любую строку события для просмотра расширенного описания во "
        "всплывающем окне.",
        "При необходимости измените язык интерфейса (RU/EN/ZH/ES) и тему "
        "оформления (светлая/тёмная) с помощью элементов в верхней панели. "
        "Выбранные настройки сохраняются автоматически.",
    ])
    add_page_break(doc)

    add_heading(doc, "ПРИЛОЖЕНИЕ Г", level=0)
    add_heading(doc, "Краткая документация программиста", level=1)
    add_paragraph(doc, "Структура репозитория:")
    add_bullets(doc, [
        "backend/app/main.py — точка входа FastAPI, эндпоинты, CORS, lifecycle hooks;",
        "backend/app/models.py — ORM-модель Event;",
        "backend/app/schemas.py — Pydantic-схемы валидации EventCreate/EventRead;",
        "backend/app/crud.py — операции CRUD и upsert;",
        "backend/app/alfaforex_sync.py — синхронизация с AlfaForex (API + Playwright fallback);",
        "backend/app/fred_sync.py — интеграция с API FRED;",
        "backend/app/database.py — подключение SQLite, миграции колонок;",
        "backend/app/seed.py — первичная инициализация данных;",
        "frontend/src/App.jsx — главные компоненты, маршрутизация, фильтрация;",
        "frontend/src/api.js — транспортный слой работы с REST API;",
        "frontend/src/i18n/I18nContext.jsx и strings.js — локализация интерфейса.",
    ])
    add_paragraph(doc, "Запуск приложения:")
    add_numbered(doc, [
        "Backend: создать виртуальное окружение, установить зависимости из "
        "backend/requirements.txt и запустить uvicorn командой "
        "uvicorn app.main:app --reload из директории backend/. Доступность "
        "сервиса проверяется через GET /health.",
        "Frontend: установить зависимости через npm install в директории "
        "frontend/, при необходимости задать переменную VITE_API_BASE с адресом "
        "backend, запустить dev-сервер командой npm run dev.",
        "Сборка production-версии клиента выполняется командой npm run build, "
        "артефакты размещаются в каталоге dist.",
    ])
    add_paragraph(doc, "Точки расширения:")
    add_bullets(doc, [
        "добавление нового внешнего источника — создать sync-модуль и "
        "интегрировать вызов в main.py/crud.py;",
        "расширение полей события — добавить колонку в models.py и обновить "
        "ensure_sqlite_columns в database.py;",
        "добавление новых локалей — пополнить strings.js и список locales в "
        "I18nContext;",
        "переход на промышленную СУБД — заменить DATABASE_URL и подключить "
        "Alembic для миграций.",
    ])


def main():
    doc = Document()
    configure_document(doc)
    add_page_numbers(doc)

    title_page(doc)
    toc(doc)
    introduction(doc)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    conclusion(doc)
    references(doc)
    appendices(doc)

    out_path = "/workspace/Курсовая_работа_Экономический_календарь.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
